import re
import language_tool_python
from docx import Document
from pathlib import Path
import argparse
import sys
import shutil

# Progress bar (tqdm) with safe fallback
try:
    from tqdm import tqdm
except Exception:
    def tqdm(iterable=None, **kwargs):
        return iterable

print("🐍 Python:", sys.version)
print("🛤️  Executable:", sys.executable)
try:
    import site, sysconfig
    print("📦 sys.path[0]:", sys.path[0])
    print("📦 site-packages:", sysconfig.get_paths().get("purelib"))
except Exception:
    pass


# CLI & path handling
parser = argparse.ArgumentParser(description="Proofread a .docx file with language_tool_python.")
parser.add_argument("--input", "-i", default="Taxation 1 QQRs - 2025.docx", help="Path to input .docx file")
parser.add_argument("--output", "-o", default="Taxation_Proofread_Output.docx", help="Path to output .docx file")
parser.add_argument("--log", "-l", default="Proofreading_Changes_Log.txt", help="Path to changes log text file")
parser.add_argument("--remote", action="store_true", help="Use LanguageTool public API instead of local Java server")
parser.add_argument("--server-url", default=None, help="Remote LanguageTool server URL (overrides --remote)")
args = parser.parse_args()

base_dir = Path(__file__).parent
input_path = Path(args.input)
if not input_path.is_absolute():
    input_path = base_dir / input_path
output_path = Path(args.output)
if not output_path.is_absolute():
    output_path = base_dir / output_path
log_path = Path(args.log)
if not log_path.is_absolute():
    log_path = base_dir / log_path

grammar_log_path = base_dir / "grammar_edits.log"
skipped_log_path = base_dir / "skipped_items.log"
no_issues_log_path = base_dir / "no_issues.log"

if not input_path.exists():
    sys.exit(f"❌ Input file not found: {input_path}")

# Work on a COPY of the original document to avoid modifying the source file
# If output path equals input path, generate a safe "_corrected" filename
if input_path.resolve() == output_path.resolve():
    output_path = input_path.with_name(input_path.stem + "_corrected" + input_path.suffix)

# Copy original to output destination and load the COPY
shutil.copy2(str(input_path), str(output_path))
doc = Document(str(output_path))

# Load grammar tool
try:
    if args.server_url:
        # Use a provided remote LanguageTool server (no local Java required)
        tool = language_tool_python.LanguageTool('en-US', remote_server=args.server_url)
    elif args.remote:
        # Use the public API (rate-limited, no local Java required)
        tool = language_tool_python.LanguageToolPublicAPI('en-US')
    else:
        # Default: start local Java-backed server
        tool = language_tool_python.LanguageTool('en-US')
except Exception as e:
    print(f"⚠️ Falling back to public API due to error initializing local/remote server: {e}")
    tool = language_tool_python.LanguageToolPublicAPI('en-US')

# Categories we allow (grammar only)
ALLOWED_CATEGORIES = {
    "GRAMMAR",
    "PUNCTUATION",
    "TYPOGRAPHY",
    "SEMANTICS"
}

def correct_grammar_only(text):
    matches = tool.check(text)
    allowed = []
    applied_rules = []

    for match in matches:
        if match.ruleIssueType.upper() in ALLOWED_CATEGORIES:
            allowed.append(match)
            # Build a concise rule summary for logging
            first_repl = match.replacements[0] if getattr(match, 'replacements', None) else ''
            cat = getattr(match, 'category', '')
            msg = getattr(match, 'message', '')
            applied_rules.append(f"{match.ruleId} [{cat}/{match.ruleIssueType}] {msg} -> '{first_repl}'")
        else:
            # Non-grammar categories are ignored silently here; we log grammar-only per user request
            pass

    corrected = language_tool_python.utils.correct(text, allowed)
    return corrected, applied_rules

# Regex patterns
citation_pattern = re.compile(r'G\.R\. No\.|G\.R\. Nos\.')
quoted_text_pattern = re.compile(r'“[^”]*”|\"[^\"]*\"|‘[^’]*’|\'[^\']*\'')


book_ref_pattern = re.compile(r'\(.*?, p\. ?\d+(-\d+)?\)')

# Extra patterns/sets to avoid false corrections
case_title_pattern = re.compile(r"\b(v\.|vs\.?|versus)\b", re.IGNORECASE)
corporate_suffixes = {"inc.", "corp.", "co.", "ltd.", "llc", "llp", "gmbh", "s.a.", "n.v."}
# A short, conservative list of obvious Tagalog words/particles (lowercase)
TAGALOG_COMMON = {
    "ang", "mga", "si", "ni", "kay", "kina", "sa", "ng", "na", "nang",
    "hindi", "wala", "meron", "dito", "doon", "iyan", "iyan", "iyon",
    "ako", "ikaw", "siya", "kami", "tayo", "kayo", "sila", "natin", "namin", "nila",
    "po", "opo", "ho", "ba", "daw", "raw", "lang", "din", "rin", "pala", "yata",
    "kasi", "pero", "kahit", "habang", "sapagkat", "dahil"
}

def _tokenize_words(s: str):
    return [w for w in re.findall(r"[A-Za-z][A-Za-z\.'-]*", s)]

def is_case_title(text: str) -> bool:
    """Detect lines like 'CIR vs. Algue, Inc.' or 'X v. Y'."""
    if not case_title_pattern.search(text):
        return False
    words = _tokenize_words(text)
    cap_words = sum(1 for w in words if w[:1].isupper())
    return cap_words >= 2

def contains_tagalog(text: str) -> bool:
    words = {w.lower() for w in _tokenize_words(text)}
    return any(w in TAGALOG_COMMON for w in words)

def is_proper_noun_heavy(text: str) -> bool:
    words = [w for w in _tokenize_words(text) if len(w) > 1]
    if not words:
        return False
    proper_like = 0
    for w in words:
        lw = w.lower()
        if w[0].isupper() or lw in corporate_suffixes or w.endswith('.') and len(w) <= 4:
            proper_like += 1
    return (proper_like / max(1, len(words))) >= 0.5

# --- Token-level protection helpers ---
from typing import List, Tuple

WORD_RE = re.compile(r"[A-Za-z][A-Za-z\.'-]*")
COMMON_SENTENCE_STARTERS = {"the", "a", "an", "in", "on", "at", "for", "of", "to", "and", "but", "or", "yes", "no"}
#
# PROTECTED_PHRASES = {"non-delegability", "bayantel", "subic", "mambisao", "mambulao", "napolcom", "miaa", "sez", "lgc"}

def _token_spans(text: str) -> List[Tuple[int, int, str]]:
    return [(m.start(), m.end(), m.group(0)) for m in WORD_RE.finditer(text)]

def _is_protected_token(tok: str) -> bool:
    t = tok
    tl = t.lower()
    # if tl in PROTECTED_PHRASES:
    #     return True
    if t.isupper() and len(t) >= 2:
        return True
    if "." in t and len(t) <= 10:
        return True
    if any(ch.isdigit() for ch in t):
        return True
    if tl.startswith("non-"):
        return True
    if t[:1].isupper() and tl not in COMMON_SENTENCE_STARTERS:
        return True
    return False

def _protected_spans(text: str) -> List[Tuple[int, int]]:
    spans = []
    for s, e, tok in _token_spans(text):
        if _is_protected_token(tok):
            spans.append((s, e))
    return spans

def _overlaps(span: Tuple[int, int], spans: List[Tuple[int, int]]) -> bool:
    s, e = span
    for ps, pe in spans:
        if not (e <= ps or s >= pe):
            return True
    return False

# --- Paragraph collection helpers ---
from typing import Iterable

def _iter_table_paragraphs(table) -> Iterable:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            # Recurse into nested tables, if any
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)

def iter_all_paragraphs(document) -> Iterable:
    """Yield all paragraphs in reading order: body, tables, headers, footers.
    Note: text inside shapes/text boxes is not exposed by python-docx and won't appear here.
    """
    # Body paragraphs first
    for p in document.paragraphs:
        yield p
    # Paragraphs inside tables in the body
    for tbl in document.tables:
        yield from _iter_table_paragraphs(tbl)
    # Headers and footers per section
    for section in document.sections:
        for p in section.header.paragraphs:
            yield p
        for tbl in section.header.tables:
            yield from _iter_table_paragraphs(tbl)
        for p in section.footer.paragraphs:
            yield p
        for tbl in section.footer.tables:
            yield from _iter_table_paragraphs(tbl)

# --- Run mapping & in-place edit helpers ---

def _runs_map(para):
    """Return (full_text, [(run, abs_start, abs_end), ...]) for a paragraph."""
    text = ""
    mapping = []
    pos = 0
    for run in para.runs:
        t = run.text
        start = pos
        end = start + len(t)
        mapping.append((run, start, end))
        text += t
        pos = end
    return text, mapping


def _apply_single_edit_to_runs(para, start, end, replacement):
    """Apply a single edit [start,end) -> replacement to paragraph runs in place.
    Works right-to-left at the caller level to keep offsets stable."""
    full_text, mapping = _runs_map(para)
    # Identify first and last affected runs
    first_i = last_i = None
    for i, (run, rs, re_) in enumerate(mapping):
        if re_ > start and rs < end:
            if first_i is None:
                first_i = i
            last_i = i
    if first_i is None:
        return None  # nothing to do
    # Edit last run tail
    last_run, ls, le = mapping[last_i]
    if last_i == first_i:
        # Edit within a single run
        local_s = start - ls
        local_e = end - ls
        last_run.text = last_run.text[:local_s] + replacement + last_run.text[local_e:]
        return last_run
    # Multiple runs: set tail on last run
    tail = last_run.text[end - ls:]
    last_run.text = tail
    # Clear middle runs entirely
    for i in range(last_i - 1, first_i, -1):
        mapping[i][0].text = ""
    # Edit head on first run and insert replacement
    first_run, fs, fe = mapping[first_i]
    head = first_run.text[:start - fs]
    first_run.text = head + replacement
    return first_run


# --- Build safe edits from LanguageTool (grammar-only + protections) ---

def build_safe_edits(text):
    matches = tool.check(text)
    protected = _protected_spans(text)
    edits = []  # (start, end, replacement, rule_str, original_slice)
    skipped_msgs = []
    applied_rules = []
    for m in matches:
        start = m.offset
        end = m.offset + m.errorLength
        if str(m.ruleIssueType).upper() not in ALLOWED_CATEGORIES:
            skipped_msgs.append(
                f"filtered_by_category @ {start}-{end}: rule={m.ruleId} cat={getattr(m,'category','')} issue={m.ruleIssueType} text='{text[start:end]}'"
            )
            continue
        if _overlaps((start, end), protected):
            skipped_msgs.append(
                f"protected_token @ {start}-{end}: rule={m.ruleId} cat={getattr(m,'category','')} issue={m.ruleIssueType} text='{text[start:end]}'"
            )
            continue
        repl = m.replacements[0] if getattr(m, 'replacements', None) else None
        if not repl:
            skipped_msgs.append(
                f"no_replacement @ {start}-{end}: rule={m.ruleId} cat={getattr(m,'category','')} issue={m.ruleIssueType} text='{text[start:end]}'"
            )
            continue
        orig_slice = text[start:end]
        edits.append((start, end, repl, f"{m.ruleId} [{getattr(m,'category','')}/{m.ruleIssueType}] {getattr(m,'message','')} -> '{repl}'", orig_slice))
        applied_rules.append(f"{m.ruleId} [{getattr(m,'category','')}/{m.ruleIssueType}] {getattr(m,'message','')} -> '{repl}'")
    # sort by start desc for in-place run editing
    edits.sort(key=lambda x: x[0], reverse=True)
    return edits, applied_rules, skipped_msgs

# Should proofread filter

def should_proofread(text):
    text_stripped = text.strip()
    if not text_stripped:
        return (False, "empty")

    # Skip citations, quoted text, book refs, and long all-caps headers
    if citation_pattern.search(text_stripped):
        return (False, "legal_citation_detected")
    if quoted_text_pattern.search(text_stripped):
        return (False, "quoted_text")
    if book_ref_pattern.search(text_stripped):
        return (False, "book_reference")
    if text_stripped.isupper() and len(text_stripped.split()) > 2:
        return (False, "all_caps_header")

    # Skip case titles and lines that are mostly proper nouns/corporate names
    if is_case_title(text_stripped):
        return (False, "case_title")
    if is_proper_noun_heavy(text_stripped):
        return (False, "proper_noun_heavy")

    # Skip lines containing obvious Tagalog particles/words
    if contains_tagalog(text_stripped):
        return (False, "tagalog_detected")

    return (True, None)

# Prepare separate logs (overwrite at start)
open(grammar_log_path, "w", encoding="utf-8").close()
open(skipped_log_path, "w", encoding="utf-8").close()
open(no_issues_log_path, "w", encoding="utf-8").close()

change_logs = []
edit_count = 0
skipped_count = 0
no_issue_count = 0

# Process each paragraph (in place; preserve layout, runs, images)
all_paragraphs = list(iter_all_paragraphs(doc))
for para in tqdm(all_paragraphs, desc="Proofreading", unit="para"):
    # Use full paragraph text (including spaces) for precise offsets
    full_text = "".join(run.text for run in para.runs)
    text_for_checks = full_text.strip()

    if full_text == "":
        continue

    should, reason = should_proofread(text_for_checks)
    if should:
        edits, applied_rules, skipped_msgs = build_safe_edits(full_text)
        # Apply edits right-to-left to keep offsets stable
        if edits:
            original = full_text
            for s, e, repl, rule_str, orig_slice in edits:
                _apply_single_edit_to_runs(para, s, e, repl)
            # Reconstruct corrected text for logging
            corrected = "".join(run.text for run in para.runs)
            if corrected != original:
                log_entry = ["--- Grammar Edit ---",
                             f"Original: {original}",
                             f"Corrected: {corrected}"]
                if applied_rules:
                    log_entry.append("Grammar rules applied:\n- " + "\n- ".join(applied_rules))
                log_text = "\n".join(log_entry) + "\n\n"
                with open(grammar_log_path, "a", encoding="utf-8") as gfh:
                    gfh.write(log_text)
                edit_count += 1
        else:
            no_issue_count += 1
            no_issue_text = ("--- No Issues ---\n"
                             f"Text: {full_text}\n\n")
            with open(no_issues_log_path, "a", encoding="utf-8") as nfh:
                nfh.write(no_issue_text)
        # Log per-suggestion skips
        for msg in skipped_msgs:
            skip_text = ("--- Skipped ---\n"
                         "Reason: protected_token or non_grammar_rule\n"
                         f"Detail: {msg}\n\n")
            with open(skipped_log_path, "a", encoding="utf-8") as sfh:
                sfh.write(skip_text)
            skipped_count += 1
    else:
        if reason != "empty":
            skip_text = ("--- Skipped ---\n"
                         f"Reason: {reason}\n"
                         f"Text: {text_for_checks}\n\n")
            with open(skipped_log_path, "a", encoding="utf-8") as sfh:
                sfh.write(skip_text)
            skipped_count += 1

# Save outputs (preserving layout)
doc.save(str(output_path))
with open(log_path, "w", encoding="utf-8") as f:
    f.write("This run produced separate logs.\n")
    f.write(f"Grammar edits log: {grammar_log_path}\n")
    f.write(f"Skipped items log: {skipped_log_path}\n")
    f.write(f"No-issues log: {no_issues_log_path}\n")
    f.write(f"Total grammar edits applied: {edit_count}\n")
    f.write(f"Total skipped items: {skipped_count}\n")
    f.write(f"Total paragraphs checked with no grammar issues: {no_issue_count}\n")

print("✅ Proofreading complete!")
print(f"📄 Output: {output_path}")
print(f"📝 Summary: {log_path}")
print(f"🧾 Grammar edits: {grammar_log_path}")
print(f"🚫 Skipped items: {skipped_log_path}")
print(f"✏️ Total grammar edits applied: {edit_count}")